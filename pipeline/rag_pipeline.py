import fitz
from docx import Document
import openpyxl
import chromadb
import google.generativeai as genai
from chromadb.utils import embedding_functions
import os
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), "../env", ".env"))

# ── Constants ─────────────────────────────────────────
DATA_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "../data")
CHROMA_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "../chroma_store")
COLLECTION_NAME = "contract_docs"
SUPPORTED_FORMATS = {".pdf", ".docx", ".xlsx"}

#Extractors
def extract_pdf(filepath):
    """Extract text from PDF file using pymupdf"""
    doc =fitz.open(filepath)
    sections=[]
    for page_num, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            sections.append(
                {
                    "content" : text,
                    "section_type" : "page",
                    "section_id" : f"page_{page_num+1}"
                }
            )
    doc.close()
    print(f"PDF extracted : {len(sections)} pages extracted")
    return sections

def extract_docx(filepath):
    doc = Document(filepath)
    sections = []
    current_heading = "Introduction"
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # detect headings
        if para.style.name.startswith("Heading"):
            # save previous section
            if current_content:
                sections.append({
                    "content": "\n".join(current_content),
                    "section_type": "heading_section",
                    "section_id": current_heading
                })
            current_heading = text
            current_content = []
        else:
            current_content.append(text)

    # also extract tables as their own sections
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                rows.append(row_text)
        if rows:
            sections.append({
                "content": "\n".join(rows),
                "section_type": "table",
                "section_id": f"table_{i+1}"
            })

    # save last section
    if current_content:
        sections.append({
            "content": "\n".join(current_content),
            "section_type": "heading_section",
            "section_id": current_heading
        })

    print(f"✅ DOCX: {len(sections)} sections extracted")
    return sections

def extract_xlsx(filepath):
    wb = openpyxl.load_workbook(filepath, data_only=True)
    sections = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                rows.append(" | ".join(row_values))
        if rows:
            sections.append({
                "content": "\n".join(rows),
                "section_type": "sheet",
                "section_id": sheet_name
            })

    print(f"✅ XLSX: {len(sections)} sheets extracted")
    return sections

def load_all_documents(folder_path):
    SUPPORTED_FORMATS = {".pdf", ".docx", ".xlsx"}
    documents = {} 
    files = [
        f for f in os.listdir(folder_path)
        if os.path.splitext(f)[1].lower() in SUPPORTED_FORMATS
    ]
    
    print(f"📁 Found {len(files)} files in '{folder_path}'")
    for filename in files:
        filepath = os.path.join(folder_path, filename)
        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".pdf":
                documents[filename] = extract_pdf(filepath)
            elif ext == ".docx":
                documents[filename] = extract_docx(filepath)
            elif ext == ".xlsx":
                documents[filename] = extract_xlsx(filepath)
        except Exception as e:
            print(f"❌ Failed: {filename} — {e}")
    
    print(f"✅ Loaded {len(documents)} documents")
    return documents   

#-- chunker --------
def chunk_sections(sections, source_filename, chunk_size=200, overlap=50):
    """
    Chunks within each section separately.
    Sections never bleed into each other.
    """
    all_chunks = []

    for section in sections:
        content = section["content"].strip()
        if not content:
            continue

        words = content.split()

        if len(words) <= chunk_size:
            # entire section fits — keep as one chunk
            all_chunks.append({
                "content": content,
                "source": source_filename,
                "section_type": section["section_type"],
                "section_id": section["section_id"]
            })
        else:
            # section too large — split with overlap
            start = 0
            part = 0
            while start < len(words):
                end = start + chunk_size
                chunk = " ".join(words[start:end])
                all_chunks.append({
                    "content": chunk,
                    "source": source_filename,
                    "section_type": section["section_type"],
                    "section_id": f"{section['section_id']}_part_{part}"
                })
                start += chunk_size - overlap
                part += 1

    return [c for c in all_chunks if len(c["content"].strip()) > 50]


#-------- Store in ChromaDB -----
def store_in_chromadb(all_chunks,all_ids,all_metadata):
    api_key = os.environ.get("GEMINI_API_KEY")
    google_gf  = embedding_functions.GoogleGenerativeAiEmbeddingFunction(
        model_name="models/gemini-embedding-001",
        api_key=api_key
        )
    
    chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)
    try:
        chroma_client.delete_collection(COLLECTION_NAME)
    except:
        pass

    collection = chroma_client.create_collection(
        name=COLLECTION_NAME,
        embedding_function=google_gf,

    )

    batch_size = 50
    for i in range(0, len(all_chunks), batch_size):
        collection.add(
            documents=all_chunks[i:i+batch_size],
            ids = all_ids[i:i+batch_size],
            metadatas=all_metadata[i:i+batch_size]
        )

    print(f"Stored {collection.count()} chunks in ChromaDB")
    return collection

# ── Main Pipeline ─────────────────────────────────────
def run_pipeline():
    print("\n" + "=" * 60)
    print("DOCUMENT INTELLIGENCE — RAG PIPELINE")
    print("=" * 60)

    # Step 1 — Load all documents
    documents = load_all_documents(DATA_FOLDER)

    # Step 2 — Chunk all documents
    all_chunks = []
    all_ids = []
    all_metadata = []

    print("\n📄 Chunking documents...")
    for doc_name, sections in documents.items():
        chunks = chunk_sections(sections, doc_name)
        for i, chunk in enumerate(chunks):
            all_chunks.append(chunk["content"])
            all_ids.append(f"{doc_name}_chunk_{i}")
            all_metadata.append({
                "source": chunk["source"],
                "section_type": chunk["section_type"],
                "section_id": chunk["section_id"],
                "format": os.path.splitext(doc_name)[1].lower()
            })
        print(f"  {doc_name}: {len(chunks)} chunks")

    print(f"\n✅ Total chunks: {len(all_chunks)}")

    # Step 3 — Store in ChromaDB
    print("\n💾 Storing in ChromaDB...")
    store_in_chromadb(all_chunks, all_ids, all_metadata)

    print("\n🎉 Pipeline complete! ChromaDB ready for agent queries.")

if __name__ == "__main__":
    run_pipeline()